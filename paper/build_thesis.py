#!/usr/bin/env python
"""인제대학교 대학원 학위논문 생성기 → paper/thesis.docx

규정: CLAUDE.md '학위논문 작성 규정' 절. 19x26cm, 신명조 10.5pt, 줄간격 200%,
전문 면번호 로마 소문자 / 본문 아라비아, 항목번호 예시B(I. 1. A. 1)).
미확정 정보는 PLACE 딕셔너리의 ○○○ 유지.

실행:  uv run python paper/build_thesis.py
"""

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

PLACE = {
    "degree_kr": "공학석사",
    "title_kr": "사족보행 로봇의 험지 보행을 위한\n강화학습 알고리즘 비교에 관한 연구",
    "title_en": "A Comparative Study of Reinforcement Learning\nAlgorithms for Rough-Terrain Locomotion\nof Quadruped Robots",
    "school_kr": "인제대학교 대학원",
    "school_en": "Graduate School, Inje University",
    "dept_kr": "○○학과 ○○학 전공",
    "dept_en": "Department of ○○○",
    "author_kr": "○　　○　　○",
    "author_en": "○○○ ○○○",
    "advisor_kr": "○　　○　　○",
    "advisor_en": "Prof. ○ ○ ○",
    "date_kr": "○○○○년 12월(또는 6월)",
    "date_en": "Dec.(또는 Jun.) 20XX",
}

BODY_FONT = "신명조"


def set_font(run, size=10.5, bold=False, font=BODY_FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._element.get_or_add_rPr()
    rf = r.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        r.append(rf)
    rf.set(qn("w:eastAsia"), font)


def para(doc, text="", size=10.5, bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         spacing=2.0, before=0, after=0, indent=None, font=BODY_FONT):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = spacing
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if indent:
        pf.first_line_indent = Cm(indent)
    for i, chunk in enumerate(text.split("\n")):
        if i:
            p.add_run().add_break()
        set_font(p.add_run(chunk), size=size, bold=bold, font=font)
    return p


def center(doc, text, size, bold=False, before=0, after=0, spacing=1.5):
    return para(doc, text, size=size, bold=bold, align=WD_ALIGN_PARAGRAPH.CENTER,
                spacing=spacing, before=before, after=after)


def setup_section(sec, numbering=None, start=None):
    """19x26cm, 여백 상3/좌3.5/우2.5/하2.5, 면번호 형식 지정."""
    sec.page_width, sec.page_height = Cm(19), Cm(26)
    sec.top_margin, sec.bottom_margin = Cm(3), Cm(2.5)
    sec.left_margin, sec.right_margin = Cm(3.5), Cm(2.5)
    sec.footer_distance = Cm(1.5)
    if numbering:
        sp = sec._sectPr
        el = sp.find(qn("w:pgNumType"))
        if el is None:
            el = OxmlElement("w:pgNumType")
            sp.append(el)
        el.set(qn("w:fmt"), numbering)
        if start is not None:
            el.set(qn("w:start"), str(start))


def footer_page_number(sec, show=True):
    sec.footer.is_linked_to_previous = False
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in list(fp.runs):
        r._element.getparent().remove(r._element)
    if not show:
        return
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    fp._p.append(fld)


def heading(doc, text, level=1):
    sizes = {1: 14, 2: 12, 3: 11}
    return para(doc, text, size=sizes.get(level, 11), bold=True,
                align=WD_ALIGN_PARAGRAPH.LEFT if level > 1 else WD_ALIGN_PARAGRAPH.CENTER,
                before=18 if level == 1 else 12, after=8)


def body(doc, text):
    return para(doc, text, indent=0.75)


# ===================================================================== 본문 내용
ABSTRACT_KR = [
"사족보행 로봇의 험지 보행 능력은 재난 현장 탐색, 산업 설비 점검, 야지 정찰 등 실제 응용의 "
"핵심 요건이나, 비정형 지형에서의 접촉 동역학은 전통적 모델 기반 제어의 한계를 드러낸다. 심층 "
"강화학습(이하 RL이라 함)은 이에 대한 유력한 대안으로 부상하였으나, 주요 RL 알고리즘들이 험지 "
"보행이라는 특정 과제에서 보이는 상대적 성능, 학습 신뢰성, 효율성에 대한 동일 조건의 체계적 "
"비교는 여전히 부족하다. 본 연구의 목표는 대표적 RL 알고리즘 6종(PPO, TRPO, A3C, SAC, TD3, "
"DDPG)을 단일한 실험 조건에서 구현·학습·평가하는 재현 가능한 비교 프레임워크를 구축하고, 지형 "
"난이도에 따른 각 알고리즘의 성능 프로파일과 그 원인을 알고리즘의 구조적 원리로부터 규명하는 "
"것이다.",
"연구 재료 및 방법으로는 GPU 병렬 물리 시뮬레이터(Isaac Lab) 상에 Unitree A1 사족보행 로봇을 "
"구현하고, 평지, 계단, 불규칙 요철 지형의 3종 지형에서 동일한 보상 함수·관측 체계·학습 예산 "
"규약 아래 각 알고리즘을 무작위 시드 3개로 반복 학습하였다. 평가는 결정론적 정책에 대하여 "
"운동 성능(전진 속도, 성공률), 안정성(전복 빈도, 자세 요동), 효율성(운반 비용), 학습 효율"
"(임계 성능 도달 표본 수·벽시계 시간)의 4개 축으로 수행하였으며, 분산분석과 Tukey 사후검정 및 "
"효과크기로 통계적 유의성을 검증하였다.",
"연구 성적으로, 평지에서는 PPO와 TRPO가 목표 속도를 충족하며 시드 간 표준편차 0.01~0.02 "
"수준의 높은 재현성을 보인 반면, 불규칙 요철 지형에서는 구도가 역전되어 SAC와 TD3가 전 시드 "
"완주에 도달하였고 특히 SAC는 시드 간 표준편차 0.002의 일관성과 최저 운반 비용을 기록하였다. "
"DDPG는 모든 조건에서 보행 획득에 실패하였으며, 탐험 방식 개선은 부분적 개선에 그치고 TD3의 "
"구조적 수정을 통해서만 회복됨을 확인하였다. 학습 효율에서는 표본 효율(off-policy 우위)과 "
"벽시계 효율(on-policy 우위)이 상반되는 결과를 얻었다.",
"이상으로부터 험지 보행을 위한 RL 알고리즘 선택은 단일 우위가 아니라 지형 난이도, 배포 병목"
"(시뮬레이션 대 실기), 재현성 요구에 따라 달라져야 하며, 지형이 거칠어질수록 최대 엔트로피 "
"기반 off-policy 계열의 수렴 신뢰성이 우위를 갖는다는 결론을 얻었다. 본 연구의 프레임워크와 "
"공개 데이터는 후속 연구의 비교 기준선으로 활용될 수 있다.",
]
KEYWORDS_KR = "사족보행 로봇, 강화학습, 험지 보행, 물리 시뮬레이션, 학습 신뢰성, 알고리즘 비교"

ABSTRACT_EN = [
"Rough-terrain locomotion is a key requirement for deploying quadruped robots in disaster "
"response, industrial inspection, and field exploration, yet contact-rich dynamics on "
"unstructured ground exposes the limits of traditional model-based control. Deep reinforcement "
"learning (RL) has emerged as a compelling alternative; however, a systematic comparison of "
"major RL algorithms under identical conditions for this specific task remains scarce. This "
"study builds a reproducible benchmarking framework that implements, trains, and evaluates six "
"representative RL algorithms (PPO, TRPO, A3C, SAC, TD3, and DDPG) under a single protocol, "
"and seeks to explain the observed performance profiles from the structural principles of each "
"algorithm.",
"A Unitree A1 quadruped was simulated in a GPU-parallel physics simulator (Isaac Lab) over "
"three terrains: flat ground, stairs, and irregular bumpy ground. Each algorithm was trained "
"with three random seeds under identical reward, observation, and budget protocols. "
"Deterministic policies were evaluated along four axes: locomotion performance (forward "
"velocity, success rate), stability (fall frequency, attitude fluctuation), efficiency (cost "
"of transport), and learning efficiency (samples and wall-clock time to threshold). "
"Statistical significance was assessed with ANOVA, Tukey HSD post-hoc tests, and effect sizes.",
"On flat ground, PPO and TRPO met the commanded velocity with near-deterministic seed "
"reproducibility. On irregular rough ground the picture inverted: SAC and TD3 completed the "
"course in every seed, with SAC showing a seed standard deviation of 0.002 and the lowest cost "
"of transport, whereas one third of on-policy seeds failed to converge within budget. DDPG "
"failed to acquire locomotion in all conditions; improved exploration yielded only partial "
"recovery, while the structural corrections of TD3 restored walking. Sample efficiency favored "
"off-policy methods by an order of magnitude, while wall-clock efficiency favored on-policy "
"methods under massive parallel simulation.",
"We conclude that algorithm selection for rough-terrain locomotion should be conditioned on "
"terrain difficulty, deployment bottleneck, and reproducibility requirements rather than on a "
"single ranking, and that maximum-entropy off-policy methods gain a reliability advantage as "
"terrain roughness increases. The framework and data are released as an open baseline for "
"subsequent studies.",
]
KEYWORDS_EN = ("quadruped robot, reinforcement learning, rough-terrain locomotion, "
               "physics simulation, training reliability, algorithm comparison")

INTRO = [
("1. 연구 과제의 의의와 중요성", [
"보행 로봇은 바퀴형 이동체가 접근할 수 없는 계단, 잔해, 요철 지면을 통과할 수 있다는 점에서 "
"재난 현장 탐색, 산업 설비 점검, 야지 운송 등 다양한 응용의 핵심 플랫폼으로 주목되어 왔다. "
"그중 사족보행 로봇은 정적 안정성과 동적 기동성의 균형이 우수하여 가장 활발히 연구되는 형태이다. "
"그러나 험지 보행은 발끝 접촉의 생성과 소멸이 반복되는 불연속 동역학, 지형과의 상호작용에서 "
"오는 큰 불확실성, 고차원 연속 행동 공간이라는 세 가지 난제를 동시에 내포한다. 전통적 모델 기반 "
"제어는 정밀한 동역학 모델과 접촉 일정의 사전 규정을 요구하므로, 모델링이 어려운 비정형 지형"
"에서는 그 성능이 급격히 저하된다1). 강화학습(이하 RL이라 함)은 시행착오 데이터로부터 제어 "
"정책을 직접 학습함으로써 이러한 모델링 부담을 우회하는 접근으로, 기초과학적으로는 접촉이 "
"많은 비선형 시스템에서의 학습 이론 검증의 장이 되고, 응용과학적으로는 실용적 보행 제어기의 "
"새로운 설계 방법론이 된다는 점에서 이중의 의의를 갖는다.",
]),
("2. 관련 지견의 분석", [
"심층 RL을 이용한 사족보행 제어는 최근 수년간 급속히 발전하였다. Hwangbo 등2)은 액추에이터 "
"신경망을 결합한 시뮬레이션 학습 정책을 실기 ANYmal 로봇에 전이하여 민첩한 운동 능력을 "
"실증하였고, Lee 등3)은 지형 정보 없이 고유수용감각만으로 험지를 주파하는 정책을, Miki 등4)은 "
"외부감각을 융합한 강인한 야지 보행을 보고하였다. Rudin 등5)은 GPU 병렬 시뮬레이션으로 수천 "
"개의 환경을 동시 구동하여 수 분 내 보행 학습이 가능함을 보였으며, 이는 본 연구가 채택한 "
"대규모 병렬 학습 체계의 근간이 된다.",
"한편 이들 성과의 대부분은 근사 정책 최적화(PPO)6) 단일 알고리즘에 기초한다. RL 알고리즘 "
"군에는 신뢰영역 기반의 TRPO7), 최대 엔트로피 기반의 SAC8), 결정론적 정책 경사 계열의 "
"DDPG10)와 그 개선인 TD39), 비동기 병렬 학습의 A3C11) 등 상이한 구조적 원리를 갖는 대안들이 "
"존재하나, 이들이 험지 보행이라는 특정 과제에서 보이는 상대적 성능에 대한 동일 조건의 비교 "
"연구는 드물다. 일반 연속 제어 벤치마크에서의 비교14)는 존재하나 보행 특유의 접촉 동역학과 "
"지형 난이도 축을 다루지 못하며, RL 연구 전반의 재현성 문제, 특히 무작위 시드에 따른 성능 "
"변동이 결론을 좌우할 수 있다는 지적15)은 시드 반복과 통계 검정을 갖춘 비교 프레임워크의 "
"필요성을 시사한다. 아울러 인간 피드백을 보상으로 통합하는 연구12,13)가 언어 모델 분야에서 "
"성과를 보이면서, 이를 로봇 보행 학습에 접목할 가능성이 열리고 있으나 이는 본 논문의 후반부 "
"심화 주제로 다룬다.",
]),
("3. 추구 내용의 유도", [
"이상의 분석으로부터 다음의 연구 요소가 유도된다. 첫째, 알고리즘 간 비교가 유의미하려면 보상 "
"함수, 관측 체계, 학습 예산, 평가 절차가 단일 규약으로 통제되어야 한다. 둘째, 험지 보행의 "
"본질을 반영하려면 지형 난이도를 독립 변수로 하는 실험 설계가 필요하다. 셋째, 시드 간 변동이 "
"큰 RL의 특성상 반복 실험과 통계적 유의성 검정, 그리고 성능 평균뿐 아니라 재현성(시드 분산) "
"자체를 1급 평가 지표로 다루어야 한다. 넷째, 표본 예산이 상이한 on-policy와 off-policy 계열의 "
"공정한 비교를 위해서는 임계 성능 도달 비용으로 정규화한 학습 효율 분석이 요구된다.",
]),
("4. 구체적 연구 주제", [
"이에 본 연구는 다음을 구체적 주제로 설정한다. GPU 병렬 시뮬레이터 상에 사족보행 로봇 "
"Unitree A1과 3종 지형(평지, 계단, 불규칙 요철)을 구현하고, RL 알고리즘 6종(PPO, TRPO, A3C, "
"SAC, TD3, DDPG)을 단일 규약으로 학습시켜, 운동 성능·안정성·효율성·학습 효율의 4축 평가 "
"체계와 통계 검정으로 지형별 성능 프로파일을 도출한다. 나아가 관찰된 성능 차이를 각 알고리즘의 "
"목적함수와 갱신 규칙의 구조적 원리로부터 해석하고, 특히 보행 획득에 실패하는 알고리즘의 실패 "
"기전을 실험적으로 분해한다.",
]),
("5. 가정 및 용어의 정의", [
"본 연구는 다음의 가정과 범위에서 수행된다. 첫째, 정책은 지형 형상 정보 없이 고유수용감각만을 "
"사용하는 맹목 보행(blind locomotion)을 가정한다. 둘째, 본 논문의 주 결과는 물리 시뮬레이션 "
"환경에서의 학습과 평가에 한정하며, 실기 전이는 후속 과제로 남긴다. 셋째, 로봇에는 전방 목표 "
"속도 1.0 m/s의 지령이 주어지며, 20초의 제한 시간 내에 시작점으로부터 5 m 이상 진행하는 것을 "
"과제 성공으로 정의한다. 이하 본문에서 근사 정책 최적화는 PPO, 신뢰영역 정책 최적화는 TRPO, "
"연성 행위자-비평자는 SAC, 쌍둥이 지연 심층 결정론적 정책 경사는 TD3, 심층 결정론적 정책 "
"경사는 DDPG, 비동기 이점 행위자-비평자는 A3C로 각각 약칭한다.",
]),
]

PURPOSE = [
("1. 궁극적 연구 목적", [
"본 연구의 궁극적 목적은 사족보행 로봇의 험지 보행 과제에서 강화학습 알고리즘의 선택 기준을 "
"실증적 근거 위에 정립하는 것이다. 즉 특정 알고리즘의 단일 우열이 아니라, 지형 난이도와 배포 "
"조건에 따라 어떤 구조적 원리가 우위를 갖는지를 규명함으로써, 향후 험지 보행 제어기 개발에서 "
"합리적 알고리즘 선택과 학습 체계 설계를 가능하게 하는 일반적 지침을 도출하고자 한다.",
]),
("2. 실행 목표", [
"위 목적을 실측 가능한 형태로 유도한 실행 목표는 다음과 같다.",
"1) 동일한 보상·관측·평가 규약 아래 RL 알고리즘 6종을 구현하고, 물리 시뮬레이션 기반의 "
"재현 가능한 험지 보행 비교 프레임워크를 구축한다.",
"2) 평지, 계단, 불규칙 요철의 3종 지형에서 알고리즘별 운동 성능, 안정성, 효율성 지표를 "
"무작위 시드 3회 반복으로 측정하고 분산분석과 사후검정으로 차이의 유의성을 검정한다.",
"3) 시드 간 성능 분산으로 정의되는 학습 재현성과, 임계 성능 도달에 요구되는 표본 수 및 "
"벽시계 시간으로 정의되는 학습 효율을 지형별로 산출하여 on-policy와 off-policy 계열을 "
"정규화 비교한다.",
"4) 보행 획득에 실패하는 알고리즘에 대하여 탐험 방식과 가치 추정 구조를 분리한 대조 실험으로 "
"실패 기전을 규명한다.",
"5) 이상의 성적을 각 알고리즘의 목적함수·갱신 규칙과 연관지어 해석하고, 지형 난이도에 따른 "
"알고리즘 선택 지침을 제시한다.",
]),
]


MATERIALS = [
("1. 연구 설계 개요", [
"본 연구는 강화학습 알고리즘 6종을 독립 변수의 한 축으로, 지형 3종을 다른 한 축으로 하는 "
"요인 설계(factorial design)를 채택하였다. 알고리즘과 지형의 각 조합에 대하여 무작위 시드를 "
"달리한 3회의 독립 학습을 수행하였으며, 학습된 정책은 학습에 사용되지 않은 평가 시드 아래 "
"결정론적으로 구동하여 관찰 항목을 측정하였다. 전체 절차는 실험 환경 구축, 알고리즘 구현 및 "
"검증, 본 학습, 평가, 통계 분석의 순으로 진행되었고, 모든 실험 조건은 계층적 설정 파일로 "
"기록되어 임의 조합의 재현이 가능하도록 하였다.",
]),
("2. 연구 재료", [
"A. 시뮬레이션 환경",
"물리 시뮬레이션에는 GPU 병렬 로봇 학습 프레임워크인 Isaac Lab 0.50.2(Isaac Sim 5.1, NVIDIA, "
"미국)를 사용하였다16). 물리 엔진은 PhysX로서 적분 시간 간격은 0.005초(200 Hz)로 하였고, 제어 "
"정책은 4스텝마다 호출되어 50 Hz로 동작하였다. 병렬 환경 수는 on-policy 계열의 학습에서 "
"4,096개, off-policy 계열에서 128개로 하였는데, 이는 두 계열의 표본 수집과 갱신 빈도의 균형이 "
"상이한 점을 고려한 것이다. 모든 연산은 단일 계산 장비(GPU：GeForce RTX 4080 16 GB, NVIDIA, "
"미국; CPU：Core i5-13400, Intel, 미국; Ubuntu 22.04)에서 수행하였다.",
"B. 로봇 모델",
"연구 대상 로봇은 소형 사족보행 로봇 Unitree A1(Unitree Robotics, 중국)으로, 질량 약 12 kg, "
"작동 관절 12개(다리당 고관절 외전·굴곡, 무릎 굴곡 각 1개), 관절 토크 한계 33.5 N·m의 "
"제원을 갖는다. 시뮬레이션에는 제조사 형상에 기반한 공식 모델을 사용하였고, 구동계는 직류 "
"모터 모델(강성 25, 감쇠 0.5의 관절 위치 비례-미분 제어)로 하였다. 정책의 출력 행동은 기본 "
"기립 자세 관절각에 0.25 배율로 더해지는 관절 위치 목표로 변환된다.",
"C. 지형",
"지형은 평지, 계단, 불규칙 요철의 3종으로 하였다. 계단은 단 높이 5 cm, 단 깊이 30 cm의 "
"피라미드형이며, 불규칙 요철 지형은 격자 높이장(height field)에 ±5 cm 범위의 균일 무작위 "
"높이를 부여하고 약 25 cm 간격으로 표본화하여 발 크기 규모의 요철이 형성되도록 생성하였다. "
"지면 마찰계수는 0.7~1.0 범위로 하였고, 지형 생성의 무작위성은 학습 시드와 함께 고정하여 "
"동일 조건의 재현이 가능하도록 하였다.",
"D. 연구 대상의 규정",
"본 연구가 상정하는 이론적 모집단은 고유수용감각만으로 험지를 보행하는 소형 사족보행 로봇의 "
"제어 정책 전체이다. 비교 대상 알고리즘은 정책 경사 계열의 대표성이 확립된 6종(PPO, TRPO, "
"A3C, SAC, TD3, DDPG)으로 하였으며, 각 조합의 학습 시드는 0, 1, 2의 3개로 고정하였다. 보행 "
"획득에 실패함이 사전 확인된 조건이라도 실패 기전 분석에 필요한 경우 제외하지 않고 성적에 "
"포함하였다.",
]),
("3. 연구 방법", [
"A. 과제 정의와 관측·행동 공간",
"과제는 전방 목표 속도 1.0 m/s의 지령 추종 보행으로 하였다. 에피소드 길이는 20초이며, 시작 "
"위치로부터 직선 거리 5 m 이상 진행하면 과제 성공으로 판정한다. 몸통이 지면에 접촉하거나 "
"기울기가 임계(중력 투영 성분 기준)를 초과하면 전복으로 판정하고 에피소드를 종료한다. 관측은 "
"몸통 선속도(3), 각속도(3), 중력 투영 벡터(3), 지령(3), 관절 위치 편차(12), 관절 속도(12), "
"직전 행동(12)의 48차원 고유수용감각으로 구성하였고, 각 성분에는 선속도 2.0, 각속도 0.25, "
"관절 속도 0.05의 정규화 배율을 적용하였다. 행동은 12차원 연속 벡터이다.",
"B. 보상 함수",
"보상은 다음 성분의 가중합으로 정의하였다. 속도 추종 보상 exp{-4(v-1.0)²}(v는 몸통 전방 "
"속도), 에너지 벌점 -2.5×10⁻⁵Στ²(τ는 관절 토크), 자세 벌점 -0.5(roll²+pitch²), 발 미끄럼 "
"벌점 -0.1, 관절 한계 벌점 -0.2, 행동 변화율 벌점 -0.01, 생존 보상 0.1, 전복 벌점 -10, "
"횡방향 속도 벌점 -1.0v_y², 회전 속도 벌점 -0.5ω_z², 그리고 유각기 지속 보상(발이 접지하는 "
"순간 공중 체류 시간과 0.5초의 차에 비례)이다. 유각기 지속 보상은 속도 추종만으로는 탐험 "
"잡음에 의존하는 퇴행적 보행이 학습되는 현상을 방지하기 위한 최소한의 보행 유도 장치로 "
"도입하였으며5), 이 구성은 모든 알고리즘에 동일하게 적용되었다.",
"C. 학습 절차",
"여섯 알고리즘은 모두 동일한 심층 신경망 구조를 사용하였다. 행위자 신경망은 은닉층 "
"512-256-128의 다층 퍼셉트론, 비평자 신경망은 512-512-256-128로 하였다. 공통 하이퍼파라미터는 "
"할인율 0.99, 학습률 3×10⁻⁴(A3C, DDPG는 1×10⁻⁴)로 하였고, on-policy 계열은 일반화 이점 "
"추정17)(λ=0.95)을 사용하였다. PPO는 비율 절단 0.2와 적응형 KL 벌점을 병용하였고6), TRPO는 "
"켤레기울기법 10회 반복과 KL 제약 0.01의 역추적 선형 탐색을7), SAC는 자동 온도 조절과 이중 "
"Q-신경망을8), TD3는 목표 정책 평활화(잡음 0.2, 절단 0.5)와 2회당 1회의 지연 정책 갱신을9), "
"DDPG는 Ornstein-Uhlenbeck 잡음과 파라미터 공간 잡음의 두 탐험 변형을10) 각각 원전에 따라 "
"구현하였다. 학습 예산은 on-policy 5×10⁷ 스텝, off-policy 5×10⁶ 스텝으로 하였는데, 이는 두 "
"계열의 표본 처리 구조 차이를 반영한 것으로 계열 간 비교는 학습 효율 분석에서 임계 도달 "
"비용으로 정규화하였다. 시간 제한에 의한 에피소드 절단은 진종결과 구분하여 가치 추정치를 "
"보상에 환류하는 방식으로 처리하였다5).",
"D. 평가 절차와 관찰 항목",
"평가는 학습 종료 시점의 정책을 결정론적(확률 정책의 평균 행동)으로 구동하여 에피소드 20회에 "
"대해 수행하였고, 평가용 무작위 시드는 학습 시드와 분리하였다. 관찰 항목은 다음과 같이 "
"정의하였다. 운동 성능：평균 전진 속도(시점-종점 변위/시간), 과제 성공률, 경로 효율(변위/"
"이동 경로 길이). 안정성：분당 전복 빈도, 자세 요동(roll·pitch 표준편차의 제곱합 제곱근), "
"접촉력 분산. 효율성：운반 비용 CoT=E/(mgd)(E는 관절 기계적 일률 |τ·ω|의 시간 적분, m은 "
"질량, g는 중력가속도, d는 이동 거리), 이동 거리당 제곱평균제곱근 토크. 학습 효율：평가 "
"속도 0.8 m/s 최초 도달까지의 환경 스텝 수와 벽시계 시간, 학습 곡선의 정규화 곡선하면적. "
"예산 내 임계 미도달 시드는 중도절단(censored)으로 처리하였다.",
"E. 자료 분석 기법",
"조합별 성적은 시드 3회의 평균과 표준편차로 요약하고 t-분포 기반 95% 신뢰구간을 병기하였다. "
"알고리즘 간 차이는 일원배치 분산분석으로 검정하고 효과크기 η²을 산출하였으며, 유의한 경우 "
"Tukey HSD 사후검정과 쌍별 Cohen’s d18)를 적용하였다. 유의수준은 0.05로 하였다. 모든 분석 "
"코드는 실험 코드와 함께 관리하여 성적 산출 과정 전체의 재현이 가능하도록 하였다.",
"F. 재현성 확보 체계",
"구현에는 Python과 PyTorch를 사용하였고, 난수 발생원은 단일 진입점에서 일괄 고정하였다. 각 "
"학습 실행은 해석이 완료된 전체 설정 파일과 학습 곡선, 정책 점검점(checkpoint)을 저장하며, "
"전체 코드와 실험 산출물은 공개 저장소를 통해 제공한다. 표와 그림을 포함한 모든 성적은 저장된 "
"원자료로부터 스크립트로 재생성된다.",
]),
]

REFERENCES = [
"Raibert MH. Legged robots that balance. Cambridge, MIT Press, 1986：1~233.",
"Hwangbo J, Lee J, Dosovitskiy A, et al. Learning agile and dynamic motor skills for legged robots. Sci Robot, 2019, 4：eaau5872.",
"Lee J, Hwangbo J, Wellhausen L, et al. Learning quadrupedal locomotion over challenging terrain. Sci Robot, 2020, 5：eabc5986.",
"Miki T, Lee J, Hwangbo J, et al. Learning robust perceptive locomotion for quadrupedal robots in the wild. Sci Robot, 2022, 7：eabk2822.",
"Rudin N, Hoeller D, Reist P, et al. Learning to walk in minutes using massively parallel deep reinforcement learning. Proc Conf Robot Learn, 2022, 164：91~100.",
"Schulman J, Wolski F, Dhariwal P, et al. Proximal policy optimization algorithms. arXiv preprint, 2017：arXiv:1707.06347.",
"Schulman J, Levine S, Abbeel P, et al. Trust region policy optimization. Proc Int Conf Mach Learn, 2015, 37：1889~1897.",
"Haarnoja T, Zhou A, Abbeel P, et al. Soft actor-critic：off-policy maximum entropy deep reinforcement learning with a stochastic actor. Proc Int Conf Mach Learn, 2018, 80：1861~1870.",
"Fujimoto S, van Hoof H, Meger D. Addressing function approximation error in actor-critic methods. Proc Int Conf Mach Learn, 2018, 80：1587~1596.",
"Lillicrap TP, Hunt JJ, Pritzel A, et al. Continuous control with deep reinforcement learning. Proc Int Conf Learn Represent, 2016.",
"Mnih V, Badia AP, Mirza M, et al. Asynchronous methods for deep reinforcement learning. Proc Int Conf Mach Learn, 2016, 48：1928~1937.",
"Christiano PF, Leike J, Brown TB, et al. Deep reinforcement learning from human preferences. Adv Neural Inf Process Syst, 2017, 30：4299~4307.",
"Ouyang L, Wu J, Jiang X, et al. Training language models to follow instructions with human feedback. Adv Neural Inf Process Syst, 2022, 35：27730~27744.",
"Duan Y, Chen X, Houthooft R, et al. Benchmarking deep reinforcement learning for continuous control. Proc Int Conf Mach Learn, 2016, 48：1329~1338.",
"Henderson P, Islam R, Bachman P, et al. Deep reinforcement learning that matters. Proc AAAI Conf Artif Intell, 2018, 32：3207~3214.",
"Mittal M, Yu C, Yu Q, et al. Orbit：a unified simulation framework for interactive robot learning environments. IEEE Robot Autom Lett, 2023, 8：3740~3747.",
"Schulman J, Moritz P, Levine S, et al. High-dimensional continuous control using generalized advantage estimation. Proc Int Conf Learn Represent, 2016.",
"Cohen J. Statistical power analysis for the behavioral sciences. 2nd ed. Hillsdale, Lawrence Erlbaum Associates, 1988：19~74.",
]

TOC = [
("국문초록", "ⅰ"), ("영문초록", "ⅲ"), ("목차", "ⅴ"),
("Ⅰ. 서론", "1"), ("Ⅱ. 연구목적", "○"), ("Ⅲ. 연구재료 및 방법", "○"),
("Ⅳ. 연구성적", "○"), ("Ⅴ. 고찰", "○"), ("Ⅵ. 결론", "○"),
("참고문헌", "○"), ("부록", "○"),
]


def build():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = BODY_FONT
    st.font.size = Pt(10.5)
    st._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)

    # ------------------------------------------------ 표지 (예시1)
    sec = doc.sections[0]
    setup_section(sec)
    footer_page_number(sec, show=False)
    center(doc, f"{PLACE['degree_kr']}학위논문", 14, bold=True, before=40)
    center(doc, PLACE["title_kr"], 22, bold=True, before=28, after=20)
    center(doc, PLACE["school_kr"], 16, bold=True, before=30)
    center(doc, PLACE["dept_kr"], 14, before=8)
    center(doc, PLACE["author_kr"], 14, before=8)
    center(doc, f"지도교수　{PLACE['advisor_kr']}", 14, before=20)
    center(doc, PLACE["date_kr"].split("년")[0] + "년", 14, before=40)

    # ------------------------------------------------ 표제지 (예시2)
    doc.add_section(WD_SECTION.NEW_PAGE)
    center(doc, PLACE["title_kr"], 22, before=60, after=20)
    center(doc, PLACE["school_kr"], 16, before=24)
    center(doc, PLACE["dept_kr"], 14, before=8)
    center(doc, PLACE["author_kr"], 14, before=8)
    center(doc, f"이 논문을 {PLACE['degree_kr']}논문으로 제출함", 16, before=26)
    center(doc, f"지도교수　{PLACE['advisor_kr']}", 14, before=10)
    center(doc, PLACE["date_kr"], 14, before=14)

    # ------------------------------------------------ 인준서 (예시3)
    doc.add_section(WD_SECTION.NEW_PAGE)
    center(doc, f"{PLACE['author_kr'].replace('　','')}의 {PLACE['degree_kr']}학위논문을 인정함.",
           21, before=50, after=24)
    for role in ["위원장", "위　원", "위　원"]:
        center(doc, f"{role}____________印", 14, before=12)
    center(doc, PLACE["school_kr"], 16, before=26)
    center(doc, PLACE["date_kr"], 14, before=12)

    # ------------------------------------------------ 전문부: 로마 소문자 면번호
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_section(sec, numbering="lowerRoman", start=1)
    footer_page_number(sec, show=True)

    # 국문초록
    center(doc, "국문초록", 14, bold=True, before=6, after=10)
    center(doc, PLACE["title_kr"], 13, bold=True, after=14)
    center(doc, PLACE["author_kr"], 11)
    center(doc, f"(지도교수 : {PLACE['advisor_kr']})", 11)
    center(doc, PLACE["dept_kr"].replace(" 전공", ""), 11)
    center(doc, PLACE["school_kr"], 11, after=12)
    for t in ABSTRACT_KR:
        body(doc, t)
    para(doc, f"Key Words：{KEYWORDS_KR}", before=14)

    # 영문초록
    doc.add_page_break()
    center(doc, "ABSTRACT", 14, bold=True, before=6, after=10)
    center(doc, PLACE["title_en"], 13, bold=True, after=14)
    center(doc, PLACE["author_en"], 11)
    center(doc, f"(Advisor：{PLACE['advisor_en']})", 11)
    center(doc, PLACE["dept_en"], 11)
    center(doc, PLACE["school_en"], 11, after=12)
    for t in ABSTRACT_EN:
        body(doc, t)
    para(doc, f"Key Words：{KEYWORDS_EN}", before=14)

    # 목차
    doc.add_page_break()
    center(doc, "목　　차", 14, bold=True, before=6, after=14)
    for item, page in TOC:
        p = para(doc, "", spacing=1.6)
        set_font(p.add_run(item), size=11)
        tab = OxmlElement("w:ptab")
        tab.set(qn("w:alignment"), "right")
        tab.set(qn("w:relativeTo"), "margin")
        tab.set(qn("w:leader"), "dot")
        r = p.add_run()
        r._element.append(tab)
        set_font(p.add_run(str(page)), size=11)

    # ------------------------------------------------ 본문: 아라비아 면번호
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_section(sec, numbering="decimal", start=1)
    footer_page_number(sec, show=True)

    heading(doc, "Ⅰ. 서　　론", 1)
    for sub, paras in INTRO:
        heading(doc, sub, 2)
        for t in paras:
            body(doc, t)

    heading(doc, "Ⅱ. 연구목적", 1)
    for sub, paras in PURPOSE:
        heading(doc, sub, 2)
        for t in paras:
            body(doc, t)

    heading(doc, "Ⅲ. 연구재료 및 방법", 1)
    for sub, paras in MATERIALS:
        heading(doc, sub, 2)
        for t in paras:
            if len(t) < 40 and t[1:3] == ". ":
                heading(doc, t, 3)
            else:
                body(doc, t)

    heading(doc, "참고문헌", 1)
    for i, ref in enumerate(REFERENCES, 1):
        p = para(doc, f"{i}. {ref}", spacing=1.6, after=4,
                 align=WD_ALIGN_PARAGRAPH.LEFT)

    out = __file__.replace("build_thesis.py", "thesis.docx")
    doc.save(out)
    print(f"saved: {out}")


if __name__ == "__main__":
    build()
